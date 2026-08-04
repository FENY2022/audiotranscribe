"""Export transcript results to local document formats."""

from __future__ import annotations

import json
from datetime import datetime
from io import BytesIO
from typing import Any

from utils.text_utils import build_srt


class ExportError(RuntimeError):
    """Raised when export generation fails."""


def export_txt(data: dict[str, Any]) -> bytes:
    """Export results as UTF-8 TXT."""
    try:
        sections = [
            "OLLAMA LOCAL TRANSCRIBER",
            f"Source file: {data.get('source_filename', 'Not specified')}",
            f"Date generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
            "RAW TRANSCRIPT",
            data.get("raw_transcript", ""),
            "",
            "CLEANED TRANSCRIPT",
            data.get("cleaned_transcript", ""),
            "",
            "SUMMARY",
            data.get("summary", ""),
            "",
            "MEETING MINUTES",
            data.get("minutes", ""),
            "",
            "ACTION ITEMS",
            data.get("action_items", ""),
        ]
        return "\n".join(sections).encode("utf-8")
    except Exception as exc:
        raise ExportError("TXT export failed.") from exc


def export_json(data: dict[str, Any]) -> bytes:
    """Export results as JSON."""
    try:
        payload = dict(data)
        payload["date_exported"] = datetime.now().isoformat(timespec="seconds")
        return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
    except Exception as exc:
        raise ExportError("JSON export failed.") from exc


def export_srt(segments: list[dict[str, Any]]) -> bytes:
    """Export timestamped segments as SRT."""
    try:
        return build_srt(segments).encode("utf-8")
    except Exception as exc:
        raise ExportError("SRT export failed.") from exc


def export_docx(data: dict[str, Any]) -> bytes:
    """Export results as a Word document."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as exc:
        raise ExportError("python-docx is not installed.") from exc

    try:
        document = Document()
        styles = document.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(10)

        title = document.add_heading("Ollama Local Transcriber", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = document.add_paragraph("Private Offline Audio and Video Transcription")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_heading("File Information", level=1)
        info = document.add_table(rows=0, cols=2)
        info.style = "Table Grid"
        fields = [
            ("Source filename", data.get("source_filename", "Not specified")),
            ("Date processed", data.get("date_processed", "Not specified")),
            ("Detected language", data.get("detected_language", "Not specified")),
            ("Ollama model", data.get("ollama_model", "Not specified")),
        ]
        for label, value in fields:
            row = info.add_row().cells
            row[0].text = label
            row[1].text = str(value or "Not specified")

        sections = [
            ("Raw Transcript", data.get("raw_transcript", "")),
            ("Cleaned Transcript", data.get("cleaned_transcript", "")),
            ("Summary", data.get("summary", "")),
            ("Meeting Minutes", data.get("minutes", "")),
            ("Action Items", data.get("action_items", "")),
        ]
        for heading, content in sections:
            document.add_heading(heading, level=1)
            for paragraph in str(content or "").split("\n"):
                document.add_paragraph(paragraph)

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()
    except Exception as exc:
        raise ExportError("Word export failed.") from exc


def export_pdf(data: dict[str, Any]) -> bytes:
    """Export results as a PDF using ReportLab."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
        from reportlab.platypus.flowables import KeepTogether
        from xml.sax.saxutils import escape
    except ImportError as exc:
        raise ExportError("reportlab is not installed.") from exc

    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.7 * inch,
            leftMargin=0.7 * inch,
            topMargin=0.7 * inch,
            bottomMargin=0.7 * inch,
            title="Ollama Local Transcriber Export",
        )
        styles = getSampleStyleSheet()
        story: list[Any] = []

        story.append(Paragraph("Ollama Local Transcriber", styles["Title"]))
        story.append(Paragraph("Private Offline Audio and Video Transcription", styles["Italic"]))
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"Date generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles["Normal"]))
        story.append(Paragraph(f"Source file: {escape(str(data.get('source_filename', 'Not specified')))}", styles["Normal"]))
        story.append(Paragraph(f"Detected language: {escape(str(data.get('detected_language', 'Not specified')))}", styles["Normal"]))
        story.append(Spacer(1, 12))

        sections = [
            ("Raw Transcript", data.get("raw_transcript", "")),
            ("Cleaned Transcript", data.get("cleaned_transcript", "")),
            ("Summary", data.get("summary", "")),
            ("Meeting Minutes", data.get("minutes", "")),
            ("Action Items", data.get("action_items", "")),
        ]
        for idx, (heading, content) in enumerate(sections):
            if idx:
                story.append(PageBreak())
            story.append(Paragraph(heading, styles["Heading1"]))
            blocks = str(content or "Not specified").split("\n")
            for block in blocks:
                safe = escape(block) if block.strip() else "&nbsp;"
                story.append(KeepTogether([Paragraph(safe, styles["BodyText"]), Spacer(1, 4)]))

        def add_page_number(canvas, document):
            canvas.saveState()
            canvas.setFont("Helvetica", 8)
            canvas.drawRightString(7.8 * inch, 0.35 * inch, f"Page {document.page}")
            canvas.restoreState()

        doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
        return buffer.getvalue()
    except Exception as exc:
        raise ExportError("PDF export failed.") from exc
